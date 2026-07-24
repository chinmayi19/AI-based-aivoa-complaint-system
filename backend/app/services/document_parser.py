"""
Lightweight document text extraction.

Per the assignment, production-grade OCR/document parsing is NOT required.
We support the common formats a customer complaint would arrive in:
PDF, DOCX, TXT and raw EML/plain-text email bodies.
"""
import io
import email
from email import policy

from pypdf import PdfReader
import docx


def extract_text(filename: str, raw_bytes: bytes) -> str:
    lower = filename.lower()

    if lower.endswith(".pdf"):
        return _extract_pdf(raw_bytes)
    if lower.endswith(".docx"):
        return _extract_docx(raw_bytes)
    if lower.endswith(".eml"):
        return _extract_eml(raw_bytes)
    # default: treat as plain text (covers .txt and pasted email text)
    try:
        return raw_bytes.decode("utf-8", errors="ignore")
    except Exception:
        return ""


def _extract_pdf(raw_bytes: bytes) -> str:
    reader = PdfReader(io.BytesIO(raw_bytes))
    parts = []
    for page in reader.pages:
        parts.append(page.extract_text() or "")
    return "\n".join(parts)


def _extract_docx(raw_bytes: bytes) -> str:
    document = docx.Document(io.BytesIO(raw_bytes))
    parts = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text for cell in row.cells))
    return "\n".join(parts)


def _extract_eml(raw_bytes: bytes) -> str:
    msg = email.message_from_bytes(raw_bytes, policy=policy.default)
    subject = msg.get("subject", "")
    sender = msg.get("from", "")
    body = ""
    if msg.is_multipart():
        for part in msg.walk():
            if part.get_content_type() == "text/plain":
                body = part.get_content()
                break
    else:
        body = msg.get_content()
    return f"From: {sender}\nSubject: {subject}\n\n{body}"
